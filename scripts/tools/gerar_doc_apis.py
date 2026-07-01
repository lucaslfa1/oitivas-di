# -*- coding: utf-8 -*-
"""
Gera um documento Word (.docx) com todas as APIs e endpoints do projeto Oitivas-DI.
Saida: Desktop do usuario -> APIs_Endpoints_Oitivas-DI.docx
"""
import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ----------------------------------------------------------------------------
# Paleta
AZUL = RGBColor(0x1F, 0x3B, 0x57)
AZUL_CLARO = RGBColor(0x2E, 0x6D, 0xA4)
CINZA = RGBColor(0x55, 0x55, 0x55)
BRANCO = RGBColor(0xFF, 0xFF, 0xFF)

METHOD_COLORS = {
    "GET": "2E7D32",      # verde
    "POST": "1565C0",     # azul
    "PUT": "EF6C00",      # laranja
    "DELETE": "C62828",   # vermelho
    "PATCH": "6A1B9A",    # roxo
    "WS": "00838F",       # ciano (websocket/signalr)
}


def shade_cell(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def set_cell_text(cell, text, bold=False, color=None, size=10, align=None, white=False):
    cell.text = ""
    p = cell.paragraphs[0]
    if align:
        p.alignment = align
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.bold = bold
    if white:
        run.font.color.rgb = BRANCO
    elif color:
        run.font.color.rgb = color


def add_method_badge(paragraph, method):
    run = paragraph.add_run(f" {method} ")
    run.font.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = BRANCO
    # sombreamento de fundo no run
    rpr = run._element.get_or_add_rPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), METHOD_COLORS.get(method, "555555"))
    rpr.append(shd)


def h_title(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(26)
    run.font.bold = True
    run.font.color.rgb = AZUL
    return p


def h_service(doc, text, subtitle=None):
    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(18)
    run.font.bold = True
    run.font.color.rgb = AZUL_CLARO
    # linha inferior
    pPr = p._p.get_or_add_pPr()
    pbdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '12')
    bottom.set(qn('w:space'), '4')
    bottom.set(qn('w:color'), '2E6DA4')
    pbdr.append(bottom)
    pPr.append(pbdr)
    if subtitle:
        ps = doc.add_paragraph()
        r = ps.add_run(subtitle)
        r.font.size = Pt(10.5)
        r.font.italic = True
        r.font.color.rgb = CINZA
    return p


def add_endpoint(doc, method, path, description, params=None, body=None, returns=None):
    # Linha do endpoint: badge + caminho monospace
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(2)
    add_method_badge(p, method)
    run = p.add_run("  " + path)
    run.font.name = 'Consolas'
    run.font.size = Pt(11.5)
    run.font.bold = True
    run.font.color.rgb = AZUL

    if description:
        pd = doc.add_paragraph()
        pd.paragraph_format.space_after = Pt(2)
        r = pd.add_run(description)
        r.font.size = Pt(10)
        r.font.color.rgb = CINZA

    def detail_line(label, value):
        pl = doc.add_paragraph()
        pl.paragraph_format.left_indent = Inches(0.25)
        pl.paragraph_format.space_after = Pt(1)
        rl = pl.add_run(f"{label}: ")
        rl.font.bold = True
        rl.font.size = Pt(9.5)
        rl.font.color.rgb = AZUL_CLARO
        rv = pl.add_run(value)
        rv.font.size = Pt(9.5)
        rv.font.name = 'Consolas'

    if params:
        detail_line("Parametros", params)
    if body:
        detail_line("Body", body)
    if returns:
        detail_line("Retorno", returns)


# ----------------------------------------------------------------------------
# Conteudo: definicao das APIs

doc = Document()

# Margens
for section in doc.sections:
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)

# Capa / Titulo
h_title(doc, "Documentacao de APIs e Endpoints")
psub = doc.add_paragraph()
r = psub.add_run("Projeto Oitivas-DI / SENTINEL")
r.font.size = Pt(13)
r.font.color.rgb = CINZA
pmeta = doc.add_paragraph()
r = pmeta.add_run("Gerado em 25/06/2026  -  repositorio: github.com/lucaslfa1/oitivas-di")
r.font.size = Pt(9)
r.font.italic = True
r.font.color.rgb = CINZA

# Resumo / visao geral
doc.add_paragraph()
pi = doc.add_paragraph()
r = pi.add_run("Visao geral")
r.font.size = Pt(14); r.font.bold = True; r.font.color.rgb = AZUL_CLARO
p = doc.add_paragraph(
    "O projeto e composto por tres servicos com endpoints HTTP:\n"
    "1. Backend (.NET 8 / ASP.NET Core - 'SinistroAPI') - API principal de transcricao, analise e laudos.\n"
    "2. SENTINEL Cortex (Python / FastAPI) - gateway de processamento de midia e inteligencia.\n"
    "3. Dashboard KPI (Python / Dash + Flask) - painel de indicadores com autenticacao."
)
p.runs[0].font.size = Pt(10.5)

# ====================================================================
# 1) BACKEND .NET
# ====================================================================
h_service(doc, "1. Backend - SinistroAPI (.NET / ASP.NET Core)",
          "Base URL: http://<host>:<porta>  |  Roteamento por atributos [Route] nos Controllers")

# AnaliseController
ph = doc.add_paragraph(); r = ph.add_run("Analise de Midias  (api/analisar)")
r.font.bold = True; r.font.size = Pt(12); r.font.color.rgb = AZUL
add_endpoint(doc, "POST", "/api/analisar/imagem",
             "Analisa imagem de vistoria veicular.",
             params="multipart/form-data",
             body="Arquivo (image/*, obrigatorio), Contexto (texto opcional)",
             returns="AnaliseResponse (laudo)")
add_endpoint(doc, "POST", "/api/analisar/video",
             "Analisa video de sinistro (suporta ate ~2GB via File API).",
             params="multipart/form-data",
             body="Arquivo (video/*, obrigatorio), Contexto, Duracao (opcionais)",
             returns="AnaliseResponse (laudo)")

# AnalisesController
ph = doc.add_paragraph(); r = ph.add_run("Analises Salvas  (api)")
r.font.bold = True; r.font.size = Pt(12); r.font.color.rgb = AZUL
add_endpoint(doc, "POST", "/api/salvar",
             "Salva uma nova analise no banco de dados.",
             body="AnaliseModel (JSON)",
             returns="OperacaoResponse (id, mensagem)")
add_endpoint(doc, "GET", "/api/analises",
             "Lista as ultimas 50 analises (ordenadas por data desc).",
             returns="Lista de AnaliseModel")
add_endpoint(doc, "GET", "/api/analises/{id}",
             "Busca uma analise por ID.",
             params="id (int, rota)",
             returns="AnaliseModel ou 404")
add_endpoint(doc, "DELETE", "/api/analises/{id}",
             "Deleta uma analise por ID.",
             params="id (int, rota)",
             returns="OperacaoResponse ou 404")

# AuthController
ph = doc.add_paragraph(); r = ph.add_run("Autenticacao  (api/Auth)")
r.font.bold = True; r.font.size = Pt(12); r.font.color.rgb = AZUL
add_endpoint(doc, "POST", "/api/Auth/login",
             "Autentica usuario (username + senha).",
             body="LoginRequest { Username, Password }",
             returns="LoginResponse { Success, Message, Username, Role } | 401")
add_endpoint(doc, "POST", "/api/Auth/register",
             "Registra novo usuario (perfil padrao 'Membro', aguarda ativacao).",
             body="LoginRequest { Username, Password }",
             returns="LoginResponse | 400 (usuario ja existe)")

# FileStorageController
ph = doc.add_paragraph(); r = ph.add_run("Armazenamento de Arquivos  (api/storage)")
r.font.bold = True; r.font.size = Pt(12); r.font.color.rgb = AZUL
add_endpoint(doc, "POST", "/api/storage/upload",
             "Faz upload de arquivo de audio para wwwroot/uploads/audio.",
             params="multipart/form-data",
             body="file (IFormFile, obrigatorio)",
             returns="{ url, fileName }")

# HealthController
ph = doc.add_paragraph(); r = ph.add_run("Health / Diagnostico  (api)")
r.font.bold = True; r.font.size = Pt(12); r.font.color.rgb = AZUL
add_endpoint(doc, "GET", "/api/health",
             "Status completo (banco de dados + processador Python). Retorna healthy/degraded/unhealthy.",
             returns="JSON de status (200 ou 503)")
add_endpoint(doc, "GET", "/api/health/live",
             "Liveness probe (Kubernetes / Cloud Run).",
             returns="{ status: 'alive' }")
add_endpoint(doc, "GET", "/api/health/ready",
             "Readiness probe - pronto quando o banco esta conectado.",
             returns="{ status: 'ready' } (200) ou 503")
add_endpoint(doc, "GET", "/health",
             "Health check agregado do ASP.NET Core (MapHealthChecks).",
             returns="Healthy / Unhealthy")

# ToolsController
ph = doc.add_paragraph(); r = ph.add_run("Ferramentas  (api/tools)")
r.font.bold = True; r.font.size = Pt(12); r.font.color.rgb = AZUL
add_endpoint(doc, "POST", "/api/tools/merge-audio",
             "Une multiplos arquivos de audio em um unico MP3 (minimo 2 arquivos).",
             params="multipart/form-data",
             body="files (IFormFileCollection, >= 2)",
             returns="Arquivo merged_audio.mp3 (audio/mpeg)")

# TranscricaoController
ph = doc.add_paragraph(); r = ph.add_run("Transcricao e Laudos  (api)")
r.font.bold = True; r.font.size = Pt(12); r.font.color.rgb = AZUL
add_endpoint(doc, "POST", "/api/transcrever",
             "Transcreve audio de oitiva (STT Azure) e dispara analise de sentimento em background.",
             params="multipart/form-data; header opcional X-Connection-Id (SignalR)",
             body="Arquivo (audio/* ou video webm/mp4, obrigatorio)",
             returns="TranscricaoResponse (texto, provedor 'azure')")
add_endpoint(doc, "POST", "/api/analisar/laudo",
             "Gera laudo pericial a partir da transcricao (Azure OpenAI).",
             body="OitivaDto { Transcricao*, Duracao, Contexto }",
             returns="AnaliseResponse (laudo, 'azure-openai')")
add_endpoint(doc, "POST", "/api/analisar/oitiva",
             "Analisa a transcricao de uma oitiva.",
             body="OitivaDto { Transcricao*, Duracao, Contexto }",
             returns="AnaliseResponse (laudo)")
add_endpoint(doc, "POST", "/api/auditar",
             "Audita a conformidade de uma transcricao com base em um roteiro.",
             body="AuditoriaDto { Transcricao*, Roteiro, Contexto }",
             returns="AnaliseResponse (auditoria, 'azure-gpt4o')")
add_endpoint(doc, "POST", "/api/extrair-dados",
             "Extrai dados estruturados (JSON) da transcricao.",
             body="OitivaDto { Transcricao* }",
             returns="JSON com dados extraidos")

# SignalR Hub
ph = doc.add_paragraph(); r = ph.add_run("Tempo Real - SignalR Hub")
r.font.bold = True; r.font.size = Pt(12); r.font.color.rgb = AZUL
add_endpoint(doc, "WS", "/hubs/analysis",
             "Hub SignalR para eventos de analise em tempo real.",
             params="Metodos do cliente: JoinGroup(connectionId), SendMembersMessage(username, message)",
             returns="Evento emitido pelo servidor: ReceiveMembersMessage(username, message, timestamp)")

# ====================================================================
# 2) SENTINEL CORTEX (FastAPI)
# ====================================================================
h_service(doc, "2. SENTINEL Cortex (Python / FastAPI)",
          "Gateway de processamento de midia. Rotas disponiveis em /api/v1/... e tambem na raiz (legado). "
          "Documentacao interativa: /docs (Swagger), /redoc, /openapi.json")

# Auth
ph = doc.add_paragraph(); r = ph.add_run("Authentication")
r.font.bold = True; r.font.size = Pt(12); r.font.color.rgb = AZUL
add_endpoint(doc, "POST", "/api/v1/login",
             "Autenticacao OAuth2 (password flow), retorna token JWT.",
             params="application/x-www-form-urlencoded (username, password)",
             returns="Token { access_token, token_type: 'bearer', username }")

# Media
ph = doc.add_paragraph(); r = ph.add_run("Media Processing")
r.font.bold = True; r.font.size = Pt(12); r.font.color.rgb = AZUL
add_endpoint(doc, "POST", "/api/v1/process/audio",
             "Processa audio: normaliza volume, reduz ruido, converte para formato otimizado.",
             params="multipart/form-data",
             body="file (audio, obrigatorio), use_cache (bool, default True)",
             returns="ProcessedAudioResponse (audio base64, metadata, score de qualidade)")
add_endpoint(doc, "POST", "/api/v1/process/merge-audio",
             "Une multiplos arquivos de audio em um unico MP3 (retorno binario direto).",
             params="multipart/form-data",
             body="files (lista de arquivos)",
             returns="audio/mpeg (merged_audio.mp3) + headers X-Original/Processed/Merged-Size")
add_endpoint(doc, "POST", "/api/v1/process/video",
             "Processa video: extrai keyframes e analisa qualidade.",
             params="multipart/form-data",
             body="file (video), extract_keyframes (bool, default True), max_keyframes (int, default 10)",
             returns="ProcessedVideoResponse (keyframes base64 JPEG)")
add_endpoint(doc, "POST", "/api/v1/extract/audio",
             "Extrai a trilha de audio de um arquivo de video.",
             params="multipart/form-data",
             body="file (video MP4/AVI/MOV/MKV/WEBM/...)",
             returns="ExtractedAudioResponse (audio base64, metadados do video)")
add_endpoint(doc, "POST", "/api/v1/tools/convert-to-wav",
             "Ferramenta de debug: converte qualquer audio/video para WAV e retorna o arquivo.",
             params="multipart/form-data",
             body="file (obrigatorio)",
             returns="audio/wav (download)")

# Intelligence
ph = doc.add_paragraph(); r = ph.add_run("Intelligence")
r.font.bold = True; r.font.size = Pt(12); r.font.color.rgb = AZUL
add_endpoint(doc, "POST", "/api/v1/analyze/sentiment",
             "Analisa tom de voz / sentimento do audio (metricas acusticas).",
             params="multipart/form-data",
             body="file (audio)",
             returns="SentimentResponse { classification, description, metrics }")

# System
ph = doc.add_paragraph(); r = ph.add_run("System")
r.font.bold = True; r.font.size = Pt(12); r.font.color.rgb = AZUL
add_endpoint(doc, "GET", "/api/v1/health",
             "Health check do servico.",
             returns="{ status, service, version, cache_entries }")
add_endpoint(doc, "GET", "/api/v1/cache/stats",
             "Estatisticas do cache.",
             returns="CacheStatsResponse")
add_endpoint(doc, "POST", "/api/v1/cache/clear",
             "Limpa todo o cache.",
             returns="{ cleared: <n> }")
add_endpoint(doc, "POST", "/api/v1/cache/cleanup",
             "Remove entradas expiradas do cache.",
             returns="{ removed: <n> }")

# Outras (raiz)
ph = doc.add_paragraph(); r = ph.add_run("Outros (raiz)")
r.font.bold = True; r.font.size = Pt(12); r.font.color.rgb = AZUL
add_endpoint(doc, "GET", "/converter",
             "Serve a pagina converter.html (ferramenta de conversao).",
             returns="Arquivo HTML")
add_endpoint(doc, "GET", "/docs",
             "Swagger UI (documentacao interativa). Tambem: /redoc e /openapi.json.",
             returns="Pagina de documentacao")

# ====================================================================
# 3) DASHBOARD KPI (Dash + Flask)
# ====================================================================
h_service(doc, "3. Dashboard KPI (Python / Dash + Flask)",
          "Painel de indicadores (Plotly Dash). Possui uma rota Flask de login que valida no Backend; "
          "as demais paginas sao servidas pela SPA do Dash via roteamento de URL.")
add_endpoint(doc, "POST", "/login",
             "Autentica no Backend e cria sessao. Restringe acesso por perfil (ALLOWED_DASHBOARD_ROLES).",
             params="application/x-www-form-urlencoded (username, password)",
             returns="redirect '/' (200) | 401 (credenciais) | 403 (sem permissao) | 503 (auth indisponivel)")
add_endpoint(doc, "GET", "/",
             "Pagina principal do dashboard (Dash). Exibe login ou painel conforme a sessao.",
             returns="HTML (SPA Dash)")

# Nota de rodape
doc.add_paragraph()
pn = doc.add_paragraph()
r = pn.add_run(
    "Observacao: no SENTINEL Cortex, todas as rotas listadas com prefixo /api/v1 tambem respondem na raiz "
    "(ex.: /process/audio) por compatibilidade com frontends legados. O parametro marcado com * e obrigatorio."
)
r.font.size = Pt(9); r.font.italic = True; r.font.color.rgb = CINZA

# ----------------------------------------------------------------------------
desktop = os.path.join(os.path.expanduser("~"), "Desktop")
out_path = os.path.join(desktop, "APIs_Endpoints_Oitivas-DI.docx")
doc.save(out_path)
print("OK:", out_path)
