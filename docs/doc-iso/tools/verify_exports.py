from pathlib import Path

from docx import Document
from pypdf import PdfReader


ROOT = Path(__file__).resolve().parents[1]
EXPORTS = ROOT / "exports"

FILES = {
    "markdown": EXPORTS / "Sentinel_Documentacao_ISO9001.md",
    "docx": EXPORTS / "Sentinel_Documentacao_ISO9001.docx",
    "pdf": EXPORTS / "Sentinel_Documentacao_ISO9001.pdf",
}

REQUIRED_MARKERS = [
    "# Sentinel - Documentacao Consolidada ISO 9001",
    "# 01 - Manual da Qualidade ISO 9001",
    "# 03 - Processos e Fluxos Operacionais",
    "# 04 - Fluxogramas",
    "# 06 - Riscos, Controles e Acoes Corretivas",
    "# 08 - Snippets de Codigo Auditaveis",
    "```mermaid",
    "```csharp",
    "```javascript",
    "```python",
]


def main() -> None:
    missing = [name for name, path in FILES.items() if not path.exists()]
    if missing:
        raise SystemExit(f"Arquivos ausentes: {', '.join(missing)}")

    markdown_text = FILES["markdown"].read_text(encoding="utf-8")
    missing_markers = [marker for marker in REQUIRED_MARKERS if marker not in markdown_text]
    if missing_markers:
        raise SystemExit("Marcadores ausentes no Markdown:\n- " + "\n- ".join(missing_markers))

    document = Document(FILES["docx"])
    paragraph_count = len(document.paragraphs)
    table_count = len(document.tables)
    if paragraph_count < 100:
        raise SystemExit(f"DOCX com poucos paragrafos: {paragraph_count}")

    pdf = PdfReader(str(FILES["pdf"]))
    page_count = len(pdf.pages)
    if page_count < 10:
        raise SystemExit(f"PDF com poucas paginas: {page_count}")

    print("OK - Exportacoes verificadas")
    print(f"Markdown: {FILES['markdown'].stat().st_size} bytes")
    print(f"DOCX: {FILES['docx'].stat().st_size} bytes; {paragraph_count} paragrafos; {table_count} tabelas")
    print(f"PDF: {FILES['pdf'].stat().st_size} bytes; {page_count} paginas")


if __name__ == "__main__":
    main()
