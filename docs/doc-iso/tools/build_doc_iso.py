from __future__ import annotations

import os
import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.shared import Inches, Pt, RGBColor
from reportlab.lib import colors
from reportlab.lib.enums import TA_LEFT
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import cm
from reportlab.platypus import (
    PageBreak,
    Paragraph,
    Preformatted,
    SimpleDocTemplate,
    Spacer,
    Table,
    TableStyle,
)


ROOT = Path(__file__).resolve().parents[1]
EXPORTS = ROOT / "exports"
DOCS = [
    ROOT / "00-controle-documental.md",
    ROOT / "01-manual-qualidade-iso9001.md",
    ROOT / "02-arquitetura-e-inventario.md",
    ROOT / "03-processos-fluxos.md",
    ROOT / "04-fluxogramas.md",
    ROOT / "05-matriz-rastreabilidade-iso9001.md",
    ROOT / "06-riscos-controles-acoes.md",
    ROOT / "07-validacao-evidencias.md",
    ROOT / "08-snippets-codigo.md",
    ROOT / "09-checklists-auditoria.md",
    ROOT / "templates-registros.md",
]


def read_docs() -> list[tuple[Path, str]]:
    return [(path, path.read_text(encoding="utf-8")) for path in DOCS]


def write_consolidated_markdown(docs: list[tuple[Path, str]]) -> Path:
    EXPORTS.mkdir(parents=True, exist_ok=True)
    out = EXPORTS / "Sentinel_Documentacao_ISO9001.md"
    parts = [
        "# Sentinel - Documentacao Consolidada ISO 9001\n",
        "Documento consolidado gerado a partir de `docs/doc-iso`.\n",
        "> Segredos e credenciais devem permanecer mascarados neste documento.\n",
    ]
    for path, text in docs:
        parts.append("\n---\n")
        parts.append(f"\n<!-- Fonte: {path.name} -->\n\n")
        parts.append(text.strip())
        parts.append("\n")
    out.write_text("\n".join(parts), encoding="utf-8")
    return out


def iter_blocks(markdown: str):
    lines = markdown.splitlines()
    i = 0
    while i < len(lines):
        line = lines[i]

        if line.startswith("```"):
            lang = line.strip("`").strip()
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].startswith("```"):
                code_lines.append(lines[i])
                i += 1
            yield ("code", lang, "\n".join(code_lines))
            i += 1
            continue

        if line.startswith("|") and i + 1 < len(lines) and set(lines[i + 1].replace("|", "").strip()) <= {"-", ":", " "}:
            table_lines = [line, lines[i + 1]]
            i += 2
            while i < len(lines) and lines[i].startswith("|"):
                table_lines.append(lines[i])
                i += 1
            yield ("table", table_lines)
            continue

        if line.startswith("#"):
            level = min(len(line) - len(line.lstrip("#")), 4)
            text = line[level:].strip()
            yield ("heading", level, text)
            i += 1
            continue

        if not line.strip():
            i += 1
            continue

        para = [line]
        i += 1
        while i < len(lines) and lines[i].strip() and not lines[i].startswith("#") and not lines[i].startswith("```") and not lines[i].startswith("|"):
            para.append(lines[i])
            i += 1
        yield ("paragraph", "\n".join(para))


def clean_inline(text: str) -> str:
    text = re.sub(r"`([^`]+)`", r"\1", text)
    text = re.sub(r"\*\*([^*]+)\*\*", r"\1", text)
    text = re.sub(r"\*([^*]+)\*", r"\1", text)
    text = text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
    return text


def parse_table(lines: list[str]) -> list[list[str]]:
    rows = []
    for idx, line in enumerate(lines):
        if idx == 1:
            continue
        cells = [cell.strip() for cell in line.strip().strip("|").split("|")]
        rows.append(cells)
    return rows


def add_docx_table(doc: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    max_cols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=max_cols)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for r_idx, row in enumerate(rows):
        for c_idx in range(max_cols):
            text = row[c_idx] if c_idx < len(row) else ""
            cell = table.cell(r_idx, c_idx)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            cell.text = clean_inline(text)
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(8)
                    if r_idx == 0:
                        run.bold = True
                        run.font.color.rgb = RGBColor(255, 255, 255)
                if r_idx == 0:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()


def build_docx(docs: list[tuple[Path, str]]) -> Path:
    out = EXPORTS / "Sentinel_Documentacao_ISO9001.docx"
    doc = Document()

    section = doc.sections[0]
    section.top_margin = Inches(0.6)
    section.bottom_margin = Inches(0.6)
    section.left_margin = Inches(0.65)
    section.right_margin = Inches(0.65)

    styles = doc.styles
    styles["Normal"].font.name = "Aptos"
    styles["Normal"].font.size = Pt(9)

    title = doc.add_heading("Sentinel - Documentacao ISO 9001", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = doc.add_paragraph("Documentacao consolidada para auditoria de qualidade do projeto Sentinel.")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("Base: docs/doc-iso. Segredos e credenciais devem permanecer mascarados.")
    doc.add_page_break()

    for doc_idx, (path, text) in enumerate(docs):
        if doc_idx:
            doc.add_page_break()
        for block in iter_blocks(text):
            kind = block[0]
            if kind == "heading":
                _, level, heading = block
                doc.add_heading(heading, level=min(level, 3))
            elif kind == "paragraph":
                _, para = block
                p = doc.add_paragraph()
                p.paragraph_format.space_after = Pt(4)
                run = p.add_run(clean_inline(para.replace("\n", " ")))
                run.font.size = Pt(9)
            elif kind == "code":
                _, lang, code = block
                p = doc.add_paragraph()
                p.style = "No Spacing"
                run = p.add_run(code[:3500])
                run.font.name = "Consolas"
                run.font.size = Pt(7)
                run.font.color.rgb = RGBColor(40, 40, 40)
            elif kind == "table":
                rows = parse_table(block[1])
                add_docx_table(doc, rows)

    doc.save(out)
    return out


def pdf_styles():
    styles = getSampleStyleSheet()
    styles.add(ParagraphStyle(
        name="BodySmall",
        parent=styles["BodyText"],
        fontName="Helvetica",
        fontSize=8,
        leading=10,
        alignment=TA_LEFT,
        spaceAfter=4,
    ))
    styles.add(ParagraphStyle(
        name="CodeSmall",
        parent=styles["Code"],
        fontName="Courier",
        fontSize=6,
        leading=7,
        leftIndent=6,
        rightIndent=6,
        spaceBefore=4,
        spaceAfter=6,
    ))
    styles.add(ParagraphStyle(
        name="H1Custom",
        parent=styles["Heading1"],
        fontSize=14,
        leading=17,
        textColor=colors.HexColor("#172554"),
        spaceBefore=8,
        spaceAfter=6,
    ))
    styles.add(ParagraphStyle(
        name="H2Custom",
        parent=styles["Heading2"],
        fontSize=11,
        leading=13,
        textColor=colors.HexColor("#1e3a8a"),
        spaceBefore=6,
        spaceAfter=4,
    ))
    return styles


def add_pdf_table(story: list, rows: list[list[str]], styles) -> None:
    if not rows:
        return
    max_cols = max(len(r) for r in rows)
    data = []
    for row in rows:
        cells = []
        for idx in range(max_cols):
            value = clean_inline(row[idx] if idx < len(row) else "")
            cells.append(Paragraph(value, styles["BodySmall"]))
        data.append(cells)

    page_width = A4[0] - 2.4 * cm
    col_width = page_width / max_cols
    table = Table(data, colWidths=[col_width] * max_cols, repeatRows=1)
    table.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#1f2937")),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
        ("GRID", (0, 0), (-1, -1), 0.25, colors.HexColor("#cbd5e1")),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("LEFTPADDING", (0, 0), (-1, -1), 4),
        ("RIGHTPADDING", (0, 0), (-1, -1), 4),
        ("TOPPADDING", (0, 0), (-1, -1), 3),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
    ]))
    story.append(table)
    story.append(Spacer(1, 6))


def build_pdf(docs: list[tuple[Path, str]]) -> Path:
    out = EXPORTS / "Sentinel_Documentacao_ISO9001.pdf"
    styles = pdf_styles()
    story = [
        Paragraph("Sentinel - Documentacao ISO 9001", styles["Title"]),
        Paragraph("Documentacao consolidada para auditoria de qualidade do projeto Sentinel.", styles["BodySmall"]),
        Paragraph("Base: docs/doc-iso. Segredos e credenciais devem permanecer mascarados.", styles["BodySmall"]),
        PageBreak(),
    ]

    for doc_idx, (path, text) in enumerate(docs):
        if doc_idx:
            story.append(PageBreak())
        for block in iter_blocks(text):
            kind = block[0]
            if kind == "heading":
                _, level, heading = block
                style = styles["H1Custom"] if level <= 1 else styles["H2Custom"]
                story.append(Paragraph(clean_inline(heading), style))
            elif kind == "paragraph":
                _, para = block
                story.append(Paragraph(clean_inline(para.replace("\n", " ")), styles["BodySmall"]))
            elif kind == "code":
                _, lang, code = block
                story.append(Preformatted(code[:2500], styles["CodeSmall"]))
            elif kind == "table":
                add_pdf_table(story, parse_table(block[1]), styles)

    pdf = SimpleDocTemplate(
        str(out),
        pagesize=A4,
        rightMargin=1.2 * cm,
        leftMargin=1.2 * cm,
        topMargin=1.2 * cm,
        bottomMargin=1.2 * cm,
    )
    pdf.build(story)
    return out


def main() -> None:
    docs = read_docs()
    md = write_consolidated_markdown(docs)
    docx = build_docx(docs)
    pdf = build_pdf(docs)
    print(f"Markdown: {md}")
    print(f"DOCX: {docx}")
    print(f"PDF: {pdf}")


if __name__ == "__main__":
    main()
