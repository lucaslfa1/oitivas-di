import sys
import subprocess
import os

def install(package):
    print(f"Installing {package}...")
    subprocess.check_call([sys.executable, "-m", "pip", "install", package])

try:
    import docx
except ImportError:
    install('python-docx')
    import docx

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# Add Title
title = doc.add_heading('Documentação e Treinamento Oficial\nProjeto: Auditoria / Sentinel', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_heading('1. Visão Geral do Sistema', level=1)
doc.add_paragraph('O projeto baseia-se em inteligência artificial aplicada à segurança viária e auditoria humana. Seu objetivo principal é o mapeamento, transcrição e perícia totalmente automatizada dos diálogos que ocorrem entre o Operador (BAS/Central) e o Motorista durante ligações de sinistro de cargas/veículos.')

doc.add_heading('2. O Coração de IA (Backend C#)', level=1)
doc.add_paragraph('Os módulos cruciais encontram-se na pasta Backend/Services. Toda vez que você quiser alterar ou aprender como a IA pensa no projeto, você deve focar nestas partes do código:')

doc.add_heading('2.1. AzureWhisperService & AzureFastTranscricaoService', level=2)
doc.add_paragraph(
    'O que fazem:\n'
    'Transformam áudio em texto usando os modelos acústicos (Whisper) restritos da nuvem da Microsoft Azure.\n\n'
    'Como ensina o código:\n'
    'Eles encapsulam o array de bytes do áudio em pacotes HTTP nativos da Azure. Uma chave vital no código é a chamada pedindo o "verbose_json". Ele diz para a rede neural não devolver apenas string crua, mas um JSON contendo os "timestamps" de cada milissegundo de fala das pessoas e as probabilidades analíticas matemáticas de ser silêncio ("no_speech_prob"). \n\n'
    'Mágica Oculta:\n'
    'O código passa essa saída bruta para o algoritmo "SpeakerDetectionService", um construtor de lógica dedutiva em C# que pontua as sentenças transcritas e, com base nisso e nas pausas, separa rigidamente quem é a voz do Operador e quem é a voz do Motorista na interface.'
)

doc.add_heading('2.2. DescricaoAnaliseService (O Cérebro Computacional)', level=2)
doc.add_paragraph(
    'O que faz:\n'
    'Este código é o supervisor e orquestrador primário dos Laudos Periciais de sinistro. Após receber a transcrição do áudio totalmente picotada e processada, ele formata o texto junto das diretrizes sensíveis de perícia da empresa e submete para o modelo textual criativo (AzureOpenAIService).\n\n'
    'Como ensina o código:\n'
    'Você vai notar que ele empacota instruções rígidas via os Prompts C# literais (ex: "GetPromptAuditoriaConformidade"). Ele essencialmente atrela algemas operacionais forçando o modelo GPT-4o a encarnar as funções de um "Auditor Sênior da Opentech", solicitando que o bot cruze as transcrições com o roteiro de Segurança e traga em texto se as instruções foram seguidas.'
)

doc.add_heading('2.3. AzureOpenAIService (Interação Limpa GPT-4o)', level=2)
doc.add_paragraph(
    'O que faz:\n'
    'É o Client limpo (wrapper) que bate no portal da Azure OpenAI para realizar os cálculos cognitivos multimodais.\n\n'
    'Como ensina o código:\n'
    'O ponto de aprendizado crucial e restritivo do seu código é o controle de "alucinação". Você vai ver que ele configura rigidamente "temperature = 0.0" nas requisições do Payload, esvaziando a criatividade do robô para assegurar resultados imutáveis e repetíveis toda vez, perfeitos para checagem de apólices e documentações legais. Adicionalmente, o código suporta análise de visão computacional injetando Base64 ("image_url") para analisar fotos do acidente junto ao áudio.'
)

doc.add_heading('3. Módulos Auxiliares', level=1)
doc.add_paragraph(
    '1. Pipeline de Dados Ocultos ("sentinel-cortex"): A inteligência pesada também trabalha chamando bibliotecas de extrações ricas desenvolvidas em pacotes de linguagem Python (FastAPI). Isso alivia a carga do backend de validações extensas em metadados.\n'
    '2. Scripts Locais (ex: test_*, deploy_gcp.md): Ficam no repositório com o papel de orquestrar infraestrutura de banco de dados (SQLite/PostgreSQL) ou gerar containers em Cloud Run.\n'
    '3. Chat/Dashboards (Interface): Arquivos Frontend consumem essas requisições REST da máquina do GPT-4o.'
)

doc.add_heading('4. Ciclo de Vida Prático de Dados', level=1)
doc.add_paragraph(
    '-> 1. O áudio (.wav/.mp3) do sinistro sobe para o servidor.\n'
    '-> 2. Backend ingere, corta e envia à porta FastTranscricao (Whisper).\n'
    '-> 3. Texto é extraído mantendo o eixo temporal - jargões errados sofrem purificação textual por bibliotecas de Regex.\n'
    '-> 4. O C# formata esta transcrição limpa e invoca DescricaoAnaliseService (GPT-4o).\n'
    '-> 5. GPT audita conformidade daquele funcionário, deduz emoções da voz e arranca chaves de Placa/Carreta.\n'
    '-> 6. Status final é salvo no banco de dados para os gestores lerem na web.'
)

doc.add_heading('5. Próximos Passos de Manutenção', level=1)
doc.add_paragraph(
    'Dicas Ouro para novos desenvolvedores tocando no projeto:\n'
    '-> A lógica de deduzir quem é que está falando (Diarização) dentro de "SpeakerDetectionService.cs" é heurística. Alterar um valor nas condições "If" pode inverter quem o sistema acha ser o Motorista vs Operador.\n'
    '-> Sempre cuide dos tokens de limite: requisições OpenAI batem no tamanho context window máximo, então evite logar áudios gigantescamente extensos sem segmentar em lotes e fique sempre de olho nos limites de RPM de sua conta Azure.\n'
)

doc.save(r"C:\Users\lucas.afonso\Desktop\Treinamento_Auditoria.docx")
print("DOCX successfully saved to Desktop!")
